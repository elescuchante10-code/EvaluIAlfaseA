"""
Servicios de extracción documental con soporte multimodal controlado.

Principios:
- El texto nativo sigue siendo la fuente principal.
- Las imágenes solo agregan contexto breve y complementario.
- El OCR visual se intenta únicamente cuando falta texto útil.
- Si visión falla, se degrada silenciosamente a texto.
"""
import base64
import io
import json
import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from groq import Groq
import httpx
from app.core.config import get_settings

logger = logging.getLogger(__name__)

settings = get_settings()
GROQ_API_KEY = settings.GROQ_API_KEY
VISION_MODEL = os.environ.get(
    "EVALUAI_MULTIMODAL_MODEL",
    "meta-llama/llama-4-scout-17b-16e-instruct",
)
MAX_ANALYZED_IMAGES = max(1, min(int(os.environ.get("EVALUAI_MAX_VISUAL_IMAGES", "5")), 5))
MIN_NATIVE_TEXT_WORDS = max(20, int(os.environ.get("EVALUAI_MIN_NATIVE_WORDS", "25")))
MIN_NATIVE_TEXT_CHARS = max(120, int(os.environ.get("EVALUAI_MIN_NATIVE_CHARS", "140")))

IMAGE_LABEL_PATTERN = re.compile(
    r"\b(objeto|figura|gr[aá]fico|grafico|tabla|diagrama|obra|imagen|ilustraci[oó]n|foto|fotograf[ií]a|captura|ecuaci[oó]n|f[oó]rmula)\b",
    flags=re.IGNORECASE,
)
DECORATIVE_LABEL_PATTERN = re.compile(
    r"\b(logo|icono|icon|firma|sello|cabecera|header|footer|marca de agua)\b",
    flags=re.IGNORECASE,
)

TYPE_KEYWORDS = {
    "graph": "grafica",
    "chart": "grafica",
    "curve": "grafica",
    "graf": "grafica",
    "tabla": "tabla",
    "table": "tabla",
    "objeto": "objeto",
    "artefact": "objeto",
    "obra": "obra",
    "paint": "obra",
    "pintura": "obra",
    "formula": "formula",
    "fórmula": "formula",
    "ecuacion": "formula",
    "ecuación": "formula",
    "equation": "formula",
    "diagrama": "diagrama",
    "diagram": "diagrama",
    "photo": "foto",
    "foto": "foto",
    "fotograf": "foto",
}

GROQ_API_KEY = (settings.GROQ_API_KEY or "").strip().strip('"').strip("'")
_http_client = httpx.Client(http2=False, timeout=30.0)
_groq_client = Groq(api_key=GROQ_API_KEY, http_client=_http_client)
_DOCUMENT_PROCESSING_CACHE: Dict[int, Dict[str, Any]] = {}


def _normalize_text(value: Any) -> str:
    return " ".join(str(value or "").split())


def _count_words(text: str) -> int:
    return len(re.findall(r"\S+", str(text or "")))


def _confidence_label_to_score(label: Any, default: float = 0.35) -> float:
    normalized = str(label or "").strip().lower()
    if normalized == "high":
        return 0.85
    if normalized == "medium":
        return 0.6
    if normalized == "low":
        return 0.35
    return default


def _score_to_confidence_label(score: float) -> str:
    if score >= 0.75:
        return "high"
    if score >= 0.5:
        return "medium"
    return "low"


def _normalize_source_type(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    if normalized in {"native_text", "scanned_printed", "scanned_handwritten", "mixed"}:
        return normalized
    return "scanned_printed"


def _to_probability(score: int) -> str:
    if score >= 6:
        return "high"
    if score >= 3:
        return "medium"
    return "low"


def _probable_type_from_text(text: str) -> str:
    lowered = _normalize_text(text).lower()
    for keyword, resolved_type in TYPE_KEYWORDS.items():
        if keyword in lowered:
            return resolved_type
    return "foto"


def _candidate_summary_hint(candidate: Dict[str, Any]) -> str:
    nearby_text = _normalize_text(candidate.get("nearby_text"))
    if nearby_text:
        return nearby_text[:220]
    return candidate.get("source_label") or "Sin referencia textual cercana"


def _score_visual_candidate(candidate: Dict[str, Any]) -> int:
    score = 0
    width_px = int(candidate.get("width_px") or 0)
    height_px = int(candidate.get("height_px") or 0)
    area = width_px * height_px
    nearby_text = _normalize_text(candidate.get("nearby_text"))

    if area >= 160000:
        score += 3
    elif area >= 60000:
        score += 2
    elif area >= 25000:
        score += 1
    else:
        score -= 2

    if max(width_px, height_px) >= 900:
        score += 2
    elif max(width_px, height_px) >= 500:
        score += 1

    label_hits = IMAGE_LABEL_PATTERN.findall(nearby_text)
    if label_hits:
        score += 3

    if nearby_text and len(nearby_text) >= 80:
        score += 1

    if candidate.get("is_low_text_page"):
        score += 2

    if DECORATIVE_LABEL_PATTERN.search(nearby_text):
        score -= 3

    return score


def select_relevant_visual_candidates(candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    ranked: List[Dict[str, Any]] = []
    for candidate in candidates:
        enriched = dict(candidate)
        score = _score_visual_candidate(enriched)
        enriched["relevance_score"] = score
        enriched["probable_relevance"] = _to_probability(score)
        enriched["probable_type"] = enriched.get("probable_type") or _probable_type_from_text(
            enriched.get("nearby_text") or enriched.get("source_label")
        )
        ranked.append(enriched)

    ranked.sort(
        key=lambda item: (
            int(item.get("relevance_score") or 0),
            int(item.get("width_px") or 0) * int(item.get("height_px") or 0),
        ),
        reverse=True,
    )
    return [item for item in ranked if (item.get("relevance_score") or 0) >= 3][:MAX_ANALYZED_IMAGES]


def cache_document_processing(document_id: int, payload: Dict[str, Any]) -> None:
    if not document_id:
        return
    _DOCUMENT_PROCESSING_CACHE[document_id] = json.loads(json.dumps(payload, ensure_ascii=False))


def get_cached_document_processing(document_id: int) -> Optional[Dict[str, Any]]:
    if not document_id:
        return None
    payload = _DOCUMENT_PROCESSING_CACHE.get(document_id)
    if not payload:
        return None
    return json.loads(json.dumps(payload, ensure_ascii=False))


def _safe_import_pillow():
    try:
        from PIL import Image

        return Image
    except Exception:
        return None


def _paragraphs_from_text_blob(text: str) -> List[str]:
    paragraphs: List[str] = []
    current: List[str] = []
    for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        normalized = line.strip()
        if normalized:
            current.append(normalized)
            continue
        if current:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs


def _normalize_pdf_text_preserving_newlines(value: Any) -> str:
    """Colapsa espacios horizontales pero conserva saltos de línea (párrafos/listas en PDF)."""
    s = str(value or "")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for line in s.split("\n"):
        lines.append(" ".join(line.split()))
    return "\n".join(lines)


def _merge_tiny_segments(segments: List[str], min_words: int = 32) -> List[str]:
    """
    Evita fragmentos demasiado cortos; concatena al segmento anterior cuando el total
    sigue siendo razonable.
    """
    if not segments:
        return []
    out: List[str] = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        if not out:
            out.append(seg)
            continue
        if _count_words(seg) < min_words and _count_words(out[-1]) + _count_words(seg) <= 420:
            out[-1] = f"{out[-1]} {seg}".strip()
        else:
            out.append(seg)
    return out


def _split_oversized_text_block(
    text: str,
    max_words: int = 260,
    max_chars: int = 1300,
) -> List[str]:
    """Divide bloques largos sin párrafos explícitos en segmentos más útiles para análisis."""
    if not text.strip():
        return []
    if _count_words(text) <= max_words and len(text) <= max_chars:
        return [text.strip()]
    # 1) Oraciones (conservador; no depende de NLP pesado)
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Za-zÁÉÍÓÚÑÜáéíóúñü0-9\"«»])", text.strip())
    parts = [p.strip() for p in parts if p and p.strip()]
    if len(parts) <= 1:
        # 2) Clausulas con punto y coma
        parts = re.split(r";\s+", text.strip())
        parts = [p.strip() for p in parts if p and p.strip()]
    if len(parts) <= 1:
        # 3) Cortes por longitud de palabras (último recurso)
        words = text.split()
        chunks: List[str] = []
        buf: List[str] = []
        for w in words:
            buf.append(w)
            if len(buf) >= max_words:
                chunks.append(" ".join(buf))
                buf = []
        if buf:
            chunks.append(" ".join(buf))
        return _merge_tiny_segments(chunks)

    segments: List[str] = []
    current: List[str] = []
    cw = 0
    for sent in parts:
        sw = _count_words(sent)
        if current and cw + sw > max_words:
            segments.append(" ".join(current))
            current = [sent]
            cw = sw
        else:
            current.append(sent)
            cw += sw
    if current:
        segments.append(" ".join(current))
    return _merge_tiny_segments(segments)


def segment_text_blocks(
    text_blob: str,
    *,
    source: str = "pdf",
    max_words: int = 260,
    max_chars: int = 1300,
) -> List[str]:
    """
    Segmentación conservadora para análisis: párrafos por líneas en blanco,
    más división heurística de bloques largos (PDF nativo sin estructura clara).
    """
    raw = str(text_blob or "")
    base = _paragraphs_from_text_blob(raw)
    refined: List[str] = []
    for block in base:
        if not block.strip():
            continue
        if source == "pdf":
            refined.extend(_split_oversized_text_block(block, max_words=max_words, max_chars=max_chars))
        else:
            refined.append(block.strip())
    return [p for p in refined if p.strip()]


def _native_text_stats(paragraphs: List[str]) -> Dict[str, Any]:
    joined = "\n\n".join(paragraphs)
    word_count = _count_words(joined)
    char_count = len(joined.strip())
    sufficient = word_count >= MIN_NATIVE_TEXT_WORDS or char_count >= MIN_NATIVE_TEXT_CHARS
    return {
        "word_count": word_count,
        "char_count": char_count,
        "sufficient": sufficient,
    }


def _bytes_to_data_url(image_bytes: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(image_bytes).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _data_url_to_bytes(data_url: str) -> bytes:
    if "," not in str(data_url or ""):
        raise ValueError("Invalid data URL")
    _, encoded = data_url.split(",", 1)
    return base64.b64decode(encoded)


def _extract_docx_shape_blob(doc: Any, shape: Any) -> Tuple[Optional[bytes], Optional[str]]:
    try:
        embed_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = doc.part.related_parts[embed_id]
        content_type = getattr(image_part, "content_type", None) or "image/png"
        return image_part.blob, content_type
    except Exception:
        return None, None


def _extract_docx_payload(file_content: bytes) -> Tuple[List[str], List[Dict[str, Any]]]:
    from docx import Document

    doc = Document(io.BytesIO(file_content))
    paragraphs = [_normalize_text(para.text) for para in doc.paragraphs if _normalize_text(para.text)]
    candidates: List[Dict[str, Any]] = []
    inline_shapes = list(doc.inline_shapes)
    shape_cursor = 0

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        try:
            drawing_count = len(paragraph._element.xpath(".//*[local-name()='drawing']"))
        except Exception:
            drawing_count = 0
        if drawing_count <= 0:
            continue

        nearby_parts = []
        for neighbor in range(max(0, paragraph_index - 1), min(len(doc.paragraphs), paragraph_index + 2)):
            text = _normalize_text(doc.paragraphs[neighbor].text)
            if text:
                nearby_parts.append(text)
        nearby_text = " ".join(nearby_parts)

        for _ in range(drawing_count):
            if shape_cursor >= len(inline_shapes):
                break
            shape = inline_shapes[shape_cursor]
            shape_cursor += 1
            image_bytes, mime_type = _extract_docx_shape_blob(doc, shape)
            candidates.append(
                {
                    "asset_id": f"docx-{shape_cursor}",
                    "source": "docx",
                    "page_number": None,
                    "width_px": int((shape.width or 0) / 9525) if getattr(shape, "width", None) else 0,
                    "height_px": int((shape.height or 0) / 9525) if getattr(shape, "height", None) else 0,
                    "nearby_text": nearby_text,
                    "source_label": _normalize_text(paragraph.text),
                    "probable_type": _probable_type_from_text(nearby_text),
                    "image_bytes": image_bytes,
                    "mime_type": mime_type,
                    "is_low_text_page": False,
                }
            )

    return paragraphs, candidates


def _resolve_pdf_xobjects(page: Any) -> Dict[str, Any]:
    resources = page.get("/Resources")
    if not resources:
        return {}
    if hasattr(resources, "get_object"):
        resources = resources.get_object()
    xobjects = resources.get("/XObject")
    if not xobjects:
        return {}
    if hasattr(xobjects, "get_object"):
        xobjects = xobjects.get_object()
    return xobjects or {}


def _serialize_pdf_image_stream(image_obj: Any) -> Tuple[Optional[bytes], Optional[str]]:
    filters = image_obj.get("/Filter")
    filter_names = []
    if isinstance(filters, list):
        filter_names = [str(item) for item in filters]
    elif filters is not None:
        filter_names = [str(filters)]

    raw_bytes = getattr(image_obj, "_data", None)
    decoded_bytes = None
    try:
        decoded_bytes = image_obj.get_data()
    except Exception:
        decoded_bytes = None

    if any(name == "/DCTDecode" for name in filter_names) and raw_bytes:
        return raw_bytes, "image/jpeg"
    if any(name == "/JPXDecode" for name in filter_names) and raw_bytes:
        return raw_bytes, "image/jp2"

    Image = _safe_import_pillow()
    if Image is None or decoded_bytes is None:
        return None, None

    width = int(image_obj.get("/Width", 0) or 0)
    height = int(image_obj.get("/Height", 0) or 0)
    if width <= 0 or height <= 0:
        return None, None

    color_space = str(image_obj.get("/ColorSpace", ""))
    mode = "RGB"
    if "/DeviceGray" in color_space:
        mode = "L"
    elif "/DeviceCMYK" in color_space:
        mode = "CMYK"

    try:
        image = Image.frombytes(mode, (width, height), decoded_bytes)
    except Exception:
        return None, None

    if image.mode == "CMYK":
        image = image.convert("RGB")

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue(), "image/png"


def _extract_pdf_payload(file_content: bytes) -> Tuple[List[str], List[Dict[str, Any]]]:
    import PyPDF2

    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    full_text = []
    candidates: List[Dict[str, Any]] = []

    for page_index, page in enumerate(pdf_reader.pages, start=1):
        page_text = _normalize_pdf_text_preserving_newlines(page.extract_text() or "")
        if page_text:
            full_text.append(page_text)

        xobjects = _resolve_pdf_xobjects(page)
        for asset_index, (_, image_obj) in enumerate(xobjects.items(), start=1):
            if hasattr(image_obj, "get_object"):
                image_obj = image_obj.get_object()
            if str(image_obj.get("/Subtype")) != "/Image":
                continue

            image_bytes, mime_type = _serialize_pdf_image_stream(image_obj)
            candidates.append(
                {
                    "asset_id": f"pdf-{page_index}-{asset_index}",
                    "source": "pdf",
                    "page_number": page_index,
                    "width_px": int(image_obj.get("/Width", 0) or 0),
                    "height_px": int(image_obj.get("/Height", 0) or 0),
                    "nearby_text": _normalize_text(page_text)[:600],
                    "source_label": f"Página {page_index}",
                    "probable_type": _probable_type_from_text(page_text),
                    "image_bytes": image_bytes,
                    "mime_type": mime_type,
                    "is_low_text_page": len(_normalize_text(page_text).split()) < 40,
                }
            )

    paragraphs = segment_text_blocks("\n\n".join(full_text), source="pdf")
    return paragraphs, candidates


def _extract_txt_payload(file_content: bytes) -> Tuple[List[str], List[Dict[str, Any]]]:
    text = file_content.decode("utf-8").replace("\r\n", "\n").replace("\r", "\n")
    return [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()], []


def _parse_json_payload(raw: str, fallback: Any) -> Any:
    cleaned = _normalize_text(raw)
    if not cleaned:
        return fallback
    match = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", raw or "")
    candidate = match.group(1) if match else raw
    try:
        return json.loads(candidate)
    except Exception:
        return fallback


def _vision_request_for_candidates(candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    content: List[Dict[str, Any]] = [
        {
            "type": "text",
            "text": (
                "Analiza estas imágenes educativas y responde SOLO con JSON válido. "
                "Devuelve un arreglo llamado items. Cada item debe contener: "
                "asset_id, type (grafica|tabla|objeto|obra|formula|diagrama|foto), "
                "summary (máximo 20 palabras, descriptivo, no interpretativo), "
                "probable_relevance (high|medium|low). "
                "Las imágenes complementan el texto; no inventes significado ni evaluación extensa."
            ),
        }
    ]

    for candidate in candidates:
        if not candidate.get("image_bytes") or not candidate.get("mime_type"):
            continue
        content.append(
            {
                "type": "text",
                "text": (
                    f"asset_id: {candidate['asset_id']}. "
                    f"Contexto cercano: {_candidate_summary_hint(candidate)}"
                ),
            }
        )
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": _bytes_to_data_url(candidate["image_bytes"], candidate["mime_type"])},
            }
        )

    if len(content) <= 1:
        return []

    completion = _groq_client.chat.completions.create(
        model=VISION_MODEL,
        messages=[{"role": "user", "content": content}],
        temperature=0.0,
        max_tokens=500,
        stream=False,
        timeout=20.0,
    )
    parsed = _parse_json_payload(completion.choices[0].message.content or "", fallback={})
    items = parsed.get("items") if isinstance(parsed, dict) else None
    return items if isinstance(items, list) else []


def _ocr_request_for_candidates(candidates: List[Dict[str, Any]]) -> List[str]:
    content: List[Dict[str, Any]] = [
        {
            "type": "text",
            "text": (
                "Extrae texto útil de estas imágenes SOLO cuando sea legible. "
                "Responde SOLO con JSON válido: {\"items\":[{\"asset_id\":\"...\",\"text\":\"...\"}]}. "
                "Si una imagen no contiene texto útil, usa una cadena vacía. "
                "No interpretes; prioriza transcripción breve y exacta."
            ),
        }
    ]

    for candidate in candidates:
        if not candidate.get("image_bytes") or not candidate.get("mime_type"):
            continue
        content.append({"type": "text", "text": f"asset_id: {candidate['asset_id']}"})
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": _bytes_to_data_url(candidate["image_bytes"], candidate["mime_type"])},
            }
        )

    if len(content) <= 1:
        return []

    completion = _groq_client.chat.completions.create(
        model=VISION_MODEL,
        messages=[{"role": "user", "content": content}],
        temperature=0.0,
        max_tokens=700,
        stream=False,
        timeout=20.0,
    )
    parsed = _parse_json_payload(completion.choices[0].message.content or "", fallback={})
    items = parsed.get("items") if isinstance(parsed, dict) else []
    if not isinstance(items, list):
        return []

    paragraphs: List[str] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        text = _normalize_text(item.get("text"))
        if text:
            paragraphs.append(text)
    return paragraphs


def _transcription_request_for_candidates(candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    content: List[Dict[str, Any]] = [
        {
            "type": "text",
            "text": (
                "Transcribe estas imágenes educativas de forma conservadora y responde SOLO con JSON válido. "
                "Usa el formato {\"items\":[{\"asset_id\":\"...\","
                "\"source_type\":\"scanned_printed|scanned_handwritten|mixed\","
                "\"transcribed_paragraphs\":[\"...\"],"
                "\"transcription_confidence\":\"high|medium|low\","
                "\"low_confidence_spans\":[\"...\"]}]}. "
                "Si no puedes leer una imagen, deja transcribed_paragraphs vacío y usa confidence low. "
                "No interpretes ni completes palabras dudosas."
            ),
        }
    ]

    candidate_by_id = {}
    for candidate in candidates:
        if not candidate.get("image_bytes") or not candidate.get("mime_type"):
            continue
        candidate_by_id[str(candidate["asset_id"])] = candidate
        content.append(
            {
                "type": "text",
                "text": (
                    f"asset_id: {candidate['asset_id']}. "
                    f"Contexto cercano: {_candidate_summary_hint(candidate)}"
                ),
            }
        )
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": _bytes_to_data_url(candidate["image_bytes"], candidate["mime_type"])},
            }
        )

    if len(content) <= 1:
        return []

    completion = _groq_client.chat.completions.create(
        model=VISION_MODEL,
        messages=[{"role": "user", "content": content}],
        temperature=0.0,
        max_tokens=1200,
        stream=False,
        timeout=20.0,
    )
    parsed = _parse_json_payload(completion.choices[0].message.content or "", fallback={})
    items = parsed.get("items") if isinstance(parsed, dict) else []
    if not isinstance(items, list):
        return []

    normalized_items: List[Dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        asset_id = str(item.get("asset_id") or "").strip()
        if not asset_id or asset_id not in candidate_by_id:
            continue

        raw_paragraphs = item.get("transcribed_paragraphs")
        if isinstance(raw_paragraphs, str):
            raw_paragraphs = [raw_paragraphs]
        if not isinstance(raw_paragraphs, list):
            raw_paragraphs = [item.get("text")] if item.get("text") else []

        transcribed_paragraphs = [
            _normalize_text(paragraph)
            for paragraph in raw_paragraphs
            if _normalize_text(paragraph)
        ]

        raw_low_confidence_spans = item.get("low_confidence_spans")
        if isinstance(raw_low_confidence_spans, str):
            raw_low_confidence_spans = [raw_low_confidence_spans]
        if not isinstance(raw_low_confidence_spans, list):
            raw_low_confidence_spans = []

        low_confidence_spans: List[Dict[str, Any]] = []
        for span in raw_low_confidence_spans:
            if isinstance(span, dict):
                text = _normalize_text(span.get("text") or span.get("value") or "")
            else:
                text = _normalize_text(span)
            if text:
                low_confidence_spans.append({"text": text})

        candidate = candidate_by_id[asset_id]
        normalized_items.append(
            {
                "asset_id": asset_id,
                "page_number": candidate.get("page_number"),
                "source_type": _normalize_source_type(item.get("source_type")),
                "transcribed_paragraphs": transcribed_paragraphs,
                "transcription_confidence": _score_to_confidence_label(
                    _confidence_label_to_score(item.get("transcription_confidence"))
                ),
                "low_confidence_spans": low_confidence_spans,
            }
        )

    return normalized_items


def _build_transcription_bundle(transcription_items: List[Dict[str, Any]]) -> Dict[str, Any]:
    ordered_items = sorted(
        [item for item in transcription_items if isinstance(item, dict)],
        key=lambda item: (item.get("page_number") or 0, str(item.get("asset_id") or "")),
    )

    transcribed_paragraphs: List[str] = []
    low_confidence_spans: List[Dict[str, Any]] = []
    page_map: List[Dict[str, Any]] = []
    source_votes: List[str] = []
    confidence_scores: List[float] = []

    for item in ordered_items:
        paragraphs = [
            _normalize_text(paragraph)
            for paragraph in item.get("transcribed_paragraphs") or []
            if _normalize_text(paragraph)
        ]
        confidence_label = _score_to_confidence_label(
            _confidence_label_to_score(item.get("transcription_confidence"))
        )
        confidence_scores.append(_confidence_label_to_score(confidence_label))
        source_type = _normalize_source_type(item.get("source_type"))
        if paragraphs:
            source_votes.extend([source_type] * max(1, len(paragraphs)))

        start_index = len(transcribed_paragraphs)
        transcribed_paragraphs.extend(paragraphs)
        paragraph_indexes = list(range(start_index, len(transcribed_paragraphs)))
        page_map.append(
            {
                "asset_id": item.get("asset_id"),
                "page_number": item.get("page_number"),
                "paragraph_indexes": paragraph_indexes,
                "confidence": confidence_label,
                "source_type": source_type,
            }
        )

        for span in item.get("low_confidence_spans") or []:
            if not isinstance(span, dict):
                continue
            text = _normalize_text(span.get("text"))
            if not text:
                continue
            low_confidence_spans.append(
                {
                    "asset_id": item.get("asset_id"),
                    "page_number": item.get("page_number"),
                    "text": text,
                }
            )

    if not source_votes:
        source_type = "scanned_printed"
    elif "scanned_handwritten" in source_votes and "scanned_printed" in source_votes:
        source_type = "mixed"
    elif "scanned_handwritten" in source_votes:
        source_type = "scanned_handwritten"
    else:
        source_type = "scanned_printed"

    confidence_score = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.35

    return {
        "source_type": source_type,
        "transcribed_paragraphs": transcribed_paragraphs,
        "transcription_confidence": _score_to_confidence_label(confidence_score),
        "transcription_confidence_score": round(confidence_score, 2),
        "low_confidence_spans": low_confidence_spans,
        "page_map": page_map,
    }


def transcribe_visual_candidates(candidates: List[Dict[str, Any]]) -> Dict[str, Any]:
    return _build_transcription_bundle(_transcription_request_for_candidates(candidates))


def build_chat_image_thumbnail(
    data_url: str,
    mime_type: str,
    max_side_px: int = 384,
    jpeg_quality: int = 72,
) -> Dict[str, Any]:
    """Genera una miniatura liviana de la imagen pegada en el chat.

    Pensado para viajar dentro de `chat_image_asset` y persistir como evidencia
    honesta de una captura manuscrita. Nunca devuelve la imagen original
    completa (evita inflar el payload con un data URL de varios MB).

    Fallbacks:
      - Si Pillow no está disponible o la imagen no se puede decodificar,
        se devuelve un asset mínimo sin `thumbnail_data_url` (solo metadata).
    """
    try:
        raw_bytes = _data_url_to_bytes(data_url)
    except Exception:
        return {
            "thumbnail_data_url": None,
            "mime_type": str(mime_type or "image/png"),
            "width": 0,
            "height": 0,
        }

    Image = _safe_import_pillow()
    if Image is None:
        return {
            "thumbnail_data_url": None,
            "mime_type": str(mime_type or "image/png"),
            "width": 0,
            "height": 0,
        }

    try:
        from io import BytesIO

        with Image.open(BytesIO(raw_bytes)) as img:
            img.load()
            original_width, original_height = img.size
            if img.mode not in ("RGB", "RGBA"):
                img = img.convert("RGB")

            longest = max(original_width, original_height) or 1
            if longest > max_side_px:
                scale = max_side_px / float(longest)
                new_size = (
                    max(1, int(original_width * scale)),
                    max(1, int(original_height * scale)),
                )
                img = img.resize(new_size, Image.LANCZOS)

            out_mime = "image/jpeg"
            save_kwargs: Dict[str, Any] = {"format": "JPEG", "quality": int(jpeg_quality), "optimize": True}
            if img.mode == "RGBA":
                img = img.convert("RGB")

            buffer = BytesIO()
            img.save(buffer, **save_kwargs)
            thumb_bytes = buffer.getvalue()
            thumb_data_url = _bytes_to_data_url(thumb_bytes, out_mime)

            return {
                "thumbnail_data_url": thumb_data_url,
                "mime_type": out_mime,
                "width": img.size[0],
                "height": img.size[1],
                "original_width": original_width,
                "original_height": original_height,
            }
    except Exception:
        return {
            "thumbnail_data_url": None,
            "mime_type": str(mime_type or "image/png"),
            "width": 0,
            "height": 0,
        }


def transcribe_chat_image(data_url: str, mime_type: str, filename: str = "clipboard-image") -> Dict[str, Any]:
    candidate = {
        "asset_id": f"chat-{_normalize_text(filename) or 'image'}",
        "source": "chat",
        "page_number": 1,
        "width_px": 0,
        "height_px": 0,
        "nearby_text": "",
        "source_label": filename or "Chat image",
        "probable_type": "foto",
        "image_bytes": _data_url_to_bytes(data_url),
        "mime_type": mime_type,
        "is_low_text_page": True,
    }
    return transcribe_visual_candidates([candidate])


def _merge_vision_metadata(
    relevant_candidates: List[Dict[str, Any]],
    vision_items: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    by_asset_id = {
        str(item.get("asset_id")): item for item in vision_items if isinstance(item, dict) and item.get("asset_id")
    }
    merged: List[Dict[str, Any]] = []
    for candidate in relevant_candidates:
        item = by_asset_id.get(candidate["asset_id"], {})
        resolved_type = str(item.get("type") or candidate.get("probable_type") or "foto").strip().lower()
        if resolved_type not in {"grafica", "tabla", "objeto", "obra", "formula", "diagrama", "foto"}:
            resolved_type = candidate.get("probable_type") or "foto"
        probable_relevance = str(item.get("probable_relevance") or candidate.get("probable_relevance") or "medium").strip().lower()
        if probable_relevance not in {"high", "medium", "low"}:
            probable_relevance = candidate.get("probable_relevance") or "medium"

        summary = _normalize_text(item.get("summary"))
        if not summary:
            summary = _candidate_summary_hint(candidate)

        merged.append(
            {
                "asset_id": candidate["asset_id"],
                "page_number": candidate.get("page_number"),
                "type": resolved_type,
                "summary": summary[:220],
                "probable_relevance": probable_relevance,
            }
        )
    return merged


def extract_document_payload(filename: str, file_content: bytes) -> Dict[str, Any]:
    lowered = str(filename or "").lower()
    if lowered.endswith(".docx"):
        paragraphs, candidates = _extract_docx_payload(file_content)
    elif lowered.endswith(".pdf"):
        paragraphs, candidates = _extract_pdf_payload(file_content)
    elif lowered.endswith(".txt"):
        paragraphs, candidates = _extract_txt_payload(file_content)
    else:
        raise ValueError("Formato de archivo no soportado")

    native_paragraphs = list(paragraphs)
    native_text_available = any(paragraph.strip() for paragraph in paragraphs)
    native_text_stats = _native_text_stats(paragraphs)
    native_text_sufficient = native_text_stats["sufficient"]
    relevant_candidates = select_relevant_visual_candidates(candidates)
    visual_context: List[Dict[str, Any]] = []
    vision_failed = False
    transcription_bundle = {
        "source_type": "native_text",
        "transcribed_paragraphs": [],
        "transcription_confidence": None,
        "transcription_confidence_score": None,
        "low_confidence_spans": [],
        "page_map": [],
    }

    if relevant_candidates:
        try:
            vision_items = _vision_request_for_candidates(relevant_candidates)
            visual_context = _merge_vision_metadata(relevant_candidates, vision_items)
        except Exception as exc:
            logger.warning("Visual analysis failed; continuing with text-only fallback: %s", exc)
            vision_failed = True
            visual_context = []

    should_transcribe_visuals = bool(relevant_candidates) and (
        not native_text_sufficient or any(candidate.get("is_low_text_page") for candidate in relevant_candidates)
    )
    if should_transcribe_visuals:
        try:
            transcription_bundle = transcribe_visual_candidates(relevant_candidates)
        except Exception as exc:
            logger.warning("OCR fallback failed during multimodal upload: %s", exc)
            vision_failed = True
            transcription_bundle = {
                "source_type": "scanned_printed",
                "transcribed_paragraphs": [],
                "transcription_confidence": "low",
                "transcription_confidence_score": 0.35,
                "low_confidence_spans": [],
                "page_map": [],
            }

    transcribed_paragraphs = transcription_bundle.get("transcribed_paragraphs") or []
    if not native_text_sufficient and transcribed_paragraphs:
        paragraphs = transcribed_paragraphs

    if native_text_sufficient:
        document_source_type = "mixed" if transcribed_paragraphs else "native_text"
        document_source_confidence = 0.7 if transcribed_paragraphs else 0.95
        text_source = "mixed" if transcribed_paragraphs else "native_text"
    elif transcribed_paragraphs:
        document_source_type = _normalize_source_type(transcription_bundle.get("source_type"))
        document_source_confidence = float(transcription_bundle.get("transcription_confidence_score") or 0.35)
        text_source = "ocr_transcription"
    else:
        document_source_type = "native_text" if native_text_available else "scanned_printed"
        document_source_confidence = 0.4 if native_text_available else 0.2
        text_source = "native_text" if native_text_available else "ocr_transcription"

    return {
        "paragraphs": paragraphs,
        "processing": {
            "text_source": text_source,
            "native_text_available": native_text_available,
            "native_text_sufficient": native_text_sufficient,
            "native_text_word_count": native_text_stats["word_count"],
            "native_text_char_count": native_text_stats["char_count"],
            "document_source_type": document_source_type,
            "document_source_confidence": round(document_source_confidence, 2),
            "document_source_confidence_label": _score_to_confidence_label(document_source_confidence),
            "source_type": document_source_type,
            "source_type_confidence": round(document_source_confidence, 2),
            "original_paragraphs": native_paragraphs,
            "transcribed_paragraphs": transcribed_paragraphs,
            "transcription_confidence": transcription_bundle.get("transcription_confidence"),
            "transcription_confidence_score": transcription_bundle.get("transcription_confidence_score"),
            "low_confidence_spans": transcription_bundle.get("low_confidence_spans") or [],
            "page_map": transcription_bundle.get("page_map") or [],
            "visual_context_enabled": bool(visual_context),
            "visual_context": visual_context,
            "visual_analysis": {
                "candidate_count": len(candidates),
                "analyzed_count": len(relevant_candidates),
                "relevant_count": len(visual_context),
                "vision_failed": vision_failed,
            },
        },
    }
